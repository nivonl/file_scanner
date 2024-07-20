import pandas as pd
from docx import Document
import sys
from datetime import datetime
import numpy as np
from heuristics import heuristics_dict

def check_pii_columns_with_combinations(df, heuristics_dict):
    results = []
    for title, logic in heuristics_dict.items():
        most_fit_column = df.iloc[:, np.argmax(df.apply(lambda x: logic(x).mean(), axis=0))].name
        pct_true = np.max(df.apply(lambda x: logic(x).mean(), axis=0))
        results.append((title, most_fit_column, pct_true))
    return results

def generate_report(results, report_file_path, scanned_file_name, df, heuristics_dict):
    doc = Document()
    doc.add_heading('Personal Information Report', 0)

    doc.add_heading('Meta Data', 1)

    # Add metadata section
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Report generated on: {timestamp}")
    doc.add_paragraph(f"Scanned file: {scanned_file_name}")

    doc.add_heading('Scan Results', 1)

    if results:
        doc.add_paragraph('Personal Information Identified in Columns:')
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Heuristic'
        hdr_cells[1].text = 'Most Matched Column'
        hdr_cells[2].text = 'Percentage Match'
        
        cols_found = []
        for title, col, pct_true in results:
            row_cells = table.add_row().cells
            row_cells[0].text = title
            row_cells[1].text = col
            row_cells[2].text = f"{pct_true:.1%}"
            cols_found.append(col)
            
        # Separate section for other suspect columns
        doc.add_heading('Other columns with more than 10% match', level=1)

        for col in df.columns:
            for title, logic in heuristics_dict.items():
                match_pct = logic(df[col]).mean()
                if (match_pct > 0.1) and (col not in cols_found):
                    doc.add_paragraph(f"{title}: {col} - {match_pct:.1%}")
    else:
        doc.add_paragraph('No Personal Information Identified.')

    doc.save(report_file_path)
    print(f'Report generated and saved to {report_file_path}')

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scanner_reporter.py <csv_file_path> <report_file_path>")
        sys.exit(1)
    
    csv_file_path = sys.argv[1]
    report_file_path = sys.argv[2]

    # Load data
    df = pd.read_csv(csv_file_path)

    # Perform matching
    results = check_pii_columns_with_combinations(df, heuristics_dict)

    # Generate report
    generate_report(results, report_file_path, csv_file_path, df, heuristics_dict)