import pandas as pd
from docx import Document
import sys
from datetime import datetime
import numpy as np
import sys
sys.path.append('../tech')  

from heuristics import heuristics_dict

def check_pii_columns_with_combinations(df, heuristics_dict):
    """
    Check columns for personally identifiable information (PII) based on heuristic combinations.

    This function applies a set of heuristic functions to each column in the provided DataFrame
    and returns columns that match each heuristic by at least 70%.

    Parameters:
    df (pd.DataFrame): The DataFrame containing the data to be scanned.
    heuristics_dict (dict): A dictionary where keys are heuristic names and values are functions
                            that take a DataFrame column and return a boolean Series indicating matches.

    Returns:
    list: A list of tuples containing the heuristic name, the column name, the percentage of matches,
          the number of appearances, and the number of unique values.
    """
    results = []
    for title, logic in heuristics_dict.items():
        for col in df.columns:
            match_pct = logic(df[col]).mean()
            match_count = logic(df[col]).sum()
            if match_pct >= 0.7:
                num_appearances = df[col].value_counts().sum()
                num_unique_values = df[col].nunique()
                results.append((title, col, match_pct, match_count, num_appearances, num_unique_values))
    return results

def generate_report(results, report_file_path, scanned_file_name, df, heuristics_dict):
    """
    Generate a DOCX report based on the results of the PII column check.

    This function creates a DOCX document that summarizes the results of the PII column check.
    It includes metadata such as the report generation timestamp and the name of the scanned file,
    as well as detailed scan results including columns with at least 70% match percentages and 
    other suspect columns.

    Parameters:
    results (list): A list of tuples containing the heuristic name, the column name with the highest match, 
                    the percentage of matches in that column, the number of appearances, and the number of unique values.
    report_file_path (str): The file path where the report will be saved.
    scanned_file_name (str): The name of the file that was scanned.
    df (pd.DataFrame): The DataFrame containing the data to be scanned.
    heuristics_dict (dict): A dictionary where keys are heuristic names and values are functions
                            that take a DataFrame column and return a boolean Series indicating matches.
    """
    doc = Document()
    doc.add_heading('Personal Information Report', 0)

    doc.add_heading('Meta Data', 1)

    # Add metadata section
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Report generated on: {timestamp}")
    doc.add_paragraph(f"Scanned file: {scanned_file_name}")

    doc.add_heading('Scan Results', 1)

    if results:
        doc.add_paragraph('Personal Information Identified in Columns:')
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Heuristic'
        hdr_cells[1].text = 'Matched Column'
        hdr_cells[2].text = 'Percentage Match'
        hdr_cells[3].text = 'Number of Appearances'
        hdr_cells[4].text = 'Number of Unique Values'
        
        cols_found = []
        total_matches = {title: 0 for title in heuristics_dict.keys()}
        for title, col, pct_true, count_true, num_appearances, num_unique_values in results:
            # Populate table with results
            row_cells = table.add_row().cells
            row_cells[0].text = title
            row_cells[1].text = col
            row_cells[2].text = f"{pct_true:.1%}"
            row_cells[3].text = str(num_appearances)
            row_cells[4].text = str(num_unique_values)
            cols_found.append(col)
            total_matches[title] += count_true
            
        # Separate section for other suspect columns with more than 10% match
        doc.add_heading('Other columns with more than 10% match', level=1)
        for col in df.columns:
            for title, logic in heuristics_dict.items():
                match_pct = logic(df[col]).mean()
                if (match_pct > 0.1) and (col not in cols_found):
                    doc.add_paragraph(f"{title}: {col} - {match_pct:.1%}")

        # Add total matches section
        doc.add_heading('Total Matches Summary', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_hdr_cells = summary_table.rows[0].cells
        summary_hdr_cells[0].text = 'Heuristic'
        summary_hdr_cells[1].text = 'Total Matches'
        for title, total in total_matches.items():
            summary_row_cells = summary_table.add_row().cells
            summary_row_cells[0].text = title
            summary_row_cells[1].text = str(total)

    else:
        doc.add_paragraph('No Personal Information Identified.')

    doc.save(report_file_path)
    print(f"Report generated and saved to {report_file_path}")
    

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scanner_reporter.py <csv_file_path> <report_file_path>")
        sys.exit(1)
    
    csv_file_path = sys.argv[1]
    report_file_path = sys.argv[2]

    # Load data
    df = pd.read_csv(csv_file_path)

    # Perform matching
    results = check_pii_columns_with_combinations(df, heuristics_dict)

    # Generate report
    generate_report(results, report_file_path, csv_file_path, df, heuristics_dict)