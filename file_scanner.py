import re
import pandas as pd
from docx import Document
import sys
from datetime import datetime
import numpy as np

# Define individual rule functions
def contains_credit_card(column):
    credit_card_pattern = re.compile(r'\b(?:\d[ -]*?){13,16}\b')
    return column.apply(lambda x: bool(credit_card_pattern.match(str(x))))

def contains_special_characters(column):
    special_char_pattern = re.compile(r'[^\w\s-]')
    return column.apply(lambda x: bool(special_char_pattern.search(str(x))))

def contains_mbi(column):
    mbi_pattern = re.compile(r'^[A-HJ-NP-Z0-9]{11}$')
    return column.apply(lambda x: bool(mbi_pattern.match(str(x))))

# Define Heuristics Dictionary
heuristics_dict = {
    "Credit Card Number": contains_credit_card,
    "Special Characters": contains_special_characters,
    "Medicare Beneficiary Identifier": contains_mbi
}

def check_pii_columns_with_combinations(df, heuristics_dict):
    results = []
    for title, logic in heuristics_dict.items():
        #add explenation
        most_fit_column = df.iloc[:, np.argmax(df.apply(lambda x: logic(x).mean(), axis=0))].name
        pct_true = np.max(df.apply(lambda x: logic(x).mean(), axis=0))
        results.append((title, most_fit_column, pct_true))
    return results

def generate_report(results, report_file_path, scanned_file_name):
    doc = Document()
    doc.add_heading('Personal Information Report', 0)

    # Add timestamp and scanned file name
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Report generated on: {timestamp}")
    doc.add_paragraph(f"Scanned file: {scanned_file_name}")

    if results:
        doc.add_paragraph('Personal Information Identified in Columns:')
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Heuristic'
        hdr_cells[1].text = 'Most Matched Column'
        hdr_cells[2].text = 'Percentage Match'

        for title, col, pct_true in results:
            row_cells = table.add_row().cells
            row_cells[0].text = title
            row_cells[1].text = col
            row_cells[2].text = f"{pct_true:.1%}"

        doc.add_paragraph('Other columns with more than 10% match:')
        for col in df.columns:
            for title, logic in heuristics_dict.items():
                match_pct = logic(df[col]).mean()
                if match_pct > 0.1 and col != most_fit_column:
                    doc.add_paragraph(f"{title}: {col} - {match_pct:.1%}")
    else:
        doc.add_paragraph('No Personal Information Identified.')

    doc.save(report_file_path)
    print(f'Report generated and saved to {report_file_path}')

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python pii_report.py <csv_file_path> <report_file_path>")
        sys.exit(1)
    
    csv_file_path = sys.argv[1]
    report_file_path = sys.argv[2]

    # Load data
    df = pd.read_csv(csv_file_path)

    # Perform matching
    results = check_pii_columns_with_combinations(df, heuristics_dict)

    # Generate report
    generate_report(results, report_file_path, csv_file_path)